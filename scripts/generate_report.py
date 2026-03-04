import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(12)


def add_table(doc: Document, rows: int, cols: int, data):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = str(data[r][c]) if r < len(data) and c < len(data[r]) else ""
    return table


def main() -> str:
    # Nội dung báo cáo tùy biến cho dự án phát hiện mã do ChatGPT tạo
    ten_de_tai = (
        "Nghiên cứu, xây dựng hệ thống phát hiện mã nguồn được tạo bởi ChatGPT và ứng dụng trong các học phần thực hành lập trình"
    )
    nganh = "An toàn thông tin CAND"

    tinh_cap_thiet = (
        "Trong bối cảnh các mô hình ngôn ngữ lớn hỗ trợ sinh mã nguồn ngày càng phổ biến, nguy cơ lạm dụng trong học tập và nguy cơ vi phạm bản quyền, đạo đức học thuật gia tăng. "
        "Việc phát triển hệ thống phát hiện mã nguồn được tạo bởi ChatGPT nhằm hỗ trợ giảng viên kiểm soát chất lượng đánh giá, nâng cao tính trung thực học thuật, đồng thời giúp sinh viên hiểu đúng vai trò công cụ AI. "
        "Giải pháp góp phần đảm bảo yêu cầu về bảo mật, an ninh thông tin và tuân thủ pháp luật trong môi trường giáo dục của lực lượng CAND."
    )

    muc_tieu = (
        "Nghiên cứu các phương pháp phân loại mã nguồn do AI tạo; phân tích, thiết kế và xây dựng hệ thống web gồm giao diện nộp mã và dịch vụ phân loại. "
        "Sử dụng mô hình CodeBERT làm nền tảng, cho phép nạp checkpoint huấn luyện để nâng cao độ chính xác. "
        "Tích hợp vào quy trình nộp bài thực hành, cung cấp báo cáo điểm tin cậy và gợi ý rà soát."
    )

    noi_dung = [
        (
            "Chương 1. Cơ sở lý thuyết và khảo sát",
            [
                "1.1. Tổng quan về mã do AI tạo, đặc trưng phong cách và thách thức phát hiện",
                "1.2. Mô hình CodeBERT và các hướng tiếp cận phân loại mã",
                "1.3. Khảo sát công cụ, dữ liệu tham chiếu và tiêu chí đánh giá",
                "1.4. Đề xuất kiến trúc hệ thống phát hiện phù hợp bối cảnh học phần thực hành",
            ],
        ),
        (
            "Chương 2. Phân tích, thiết kế hệ thống",
            [
                "2.1. Yêu cầu chức năng và phi chức năng",
                "2.2. Thiết kế dịch vụ phân loại (API /predict, /predict-file, /health)",
                "2.3. Thiết kế giao diện nộp mã, xem kết quả và xuất báo cáo",
                "2.4. Cấu hình mô hình: MODEL_DIR (checkpoint) hoặc MODEL_NAME (mặc định CodeBERT)",
            ],
        ),
        (
            "Chương 3. Xây dựng và triển khai",
            [
                "3.1. Cài đặt môi trường, quản lý phụ thuộc và đóng gói",
                "3.2. Xây dựng dịch vụ FastAPI và giao diện web tĩnh",
                "3.3. Thử nghiệm chức năng dự đoán và đánh giá ban đầu",
                "3.4. Hướng tích hợp vào quy trình nộp bài và CI",
            ],
        ),
    ]

    san_pham = [
        "Hệ thống web hoàn thiện cho phép nộp mã và phân loại nguồn gốc (ChatGPT/Human)",
        "API và giao diện báo cáo kết quả, có thể mở rộng xuất CSV/JSON",
        "Tài liệu hướng dẫn sử dụng và triển khai",
    ]

    tien_do = [
        ["STT", "Nội dung công việc", "Ngày hoàn thành"],
        ["1", "Hoàn thành đề cương, khảo sát, thiết kế kiến trúc", "01/09/2024 – 30/10/2024"],
        ["2", "Xây dựng mô-đun phân loại và giao diện web", "01/11/2024 – 31/01/2025"],
        ["3", "Tích hợp, thử nghiệm, hoàn thiện báo cáo", "01/02/2025 – 31/03/2025"],
    ]

    doc = Document()

    # Căn giữa tiêu đề trang đầu
    title = doc.add_paragraph()
    title_run = title.add_run("BÁO CÁO NGHIÊN CỨU")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    add_heading(doc, "1. Tên đề tài:")
    add_paragraph(doc, ten_de_tai)

    add_heading(doc, "2. Thuộc ngành/chuyên ngành:")
    add_paragraph(doc, nganh)

    add_heading(doc, "3. Tính cấp thiết:")
    add_paragraph(doc, tinh_cap_thiet)

    add_heading(doc, "4. Mục tiêu:")
    add_paragraph(doc, muc_tieu)

    add_heading(doc, "5. Nội dung:")
    for heading, items in noi_dung:
        add_paragraph(doc, heading)
        for it in items:
            add_paragraph(doc, "- " + it)

    add_heading(doc, "6. Sản phẩm/kết quả dự kiến:")
    for it in san_pham:
        add_paragraph(doc, "- " + it)

    add_heading(doc, "7. Tiến độ thực hiện:")
    add_table(doc, rows=len(tien_do), cols=len(tien_do[0]), data=tien_do)

    doc.add_paragraph("")
    place = doc.add_paragraph("Bắc Ninh, ngày ... tháng ... năm 2024")
    place.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out_dir = os.path.join(os.getcwd(), "reports")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Bao_cao_nghien_cuu_GPTSniffer.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = main()
    print(path)


